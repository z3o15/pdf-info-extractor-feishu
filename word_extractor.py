import os
import re
from docx import Document

def fix_text_format(text: str) -> str:
    """修复Word文本格式问题：断行、连字符、空格、作者分隔符"""
    if not text:
        return ""

    # 修复连字符断行 (neuro-\nscience → neuroscience)
    text = re.sub(r'(\w)-\s*\n\s*(\w)', r'\1\2', text)

    # 修复作者分隔符（空格 + | + 空格 → 中文逗号）
    text = re.sub(r'\s*\|\s*', '，', text)

    # 按行处理未完句合并
    lines = text.splitlines()
    merged_lines = []
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue  # 跳过空行

        if (i + 1 < len(lines)):
            next_line = lines[i + 1].strip()
            # 如果行尾不是句号、冒号、分号、问号、感叹号或大写字母结尾，则合并
            if not re.search(r'[.:;?!A-Z]$', line):
                line = f"{line} {next_line}"
                lines[i + 1] = ""  # 标记下一行为空
        merged_lines.append(line)

    text = "\n".join(l for l in merged_lines if l.strip())

    # 清理多余空格
    text = re.sub(r'\s+', ' ', text)

    # 恢复段落换行（句号或问号后加换行以便阅读）
    text = re.sub(r'([。！？!?])\s*', r'\1\n', text)

    # 清理首尾空格
    return text.strip()


def extract_word_info(word_path):
    """提取Word文件信息（从起始到key words之前的所有内容）"""
    try:
        if not os.path.exists(word_path):
            print(f"文件不存在: {word_path}")
            return None

        # 检查文件扩展名
        if not word_path.lower().endswith(('.docx', '.doc')):
            print(f"不支持的文件格式: {word_path}")
            return None

        # 读取Word文档
        doc = Document(word_path)
        
        # 提取所有段落文本
        full_text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text += paragraph.text + "\n"

        if not full_text.strip():
            print(f"Word文件内容为空: {word_path}")
            return None

        # ====== 提取内容（从起始到key words之前） ======
        # 查找key words的位置（支持多种格式）
        keywords_patterns = [
            r'Key words',
            r'Keywords', 
            r'KEYWORDS',
            r'关键词',
            r'key words',
            r'keywords'
        ]
        
        keywords_position = -1
        for pattern in keywords_patterns:
            match = re.search(pattern, full_text)
            if match:
                keywords_position = match.start()
                break

        # 提取内容
        if keywords_position != -1:
            # 提取从起始到key words之前的内容
            extracted_content = full_text[:keywords_position].strip()
        else:
            # 如果没有找到key words，则提取全部内容
            extracted_content = full_text.strip()

        # 修复文本格式
        extracted_content = fix_text_format(extracted_content)

        # ====== 返回结果 ======
        return {
            '简介': extracted_content,
            '摘要': extracted_content  # Word文件暂时不区分简介和摘要，都返回相同内容
        }

    except Exception as e:
        print(f"提取Word信息时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


def extract_word_pages_direct(word_path, pages_to_extract=[1]):
    """直接从Word提取指定页面的文本（简化版，Word文档没有明确的分页概念）"""
    try:
        if not os.path.exists(word_path):
            print(f"文件不存在: {word_path}")
            return None

        # 读取Word文档
        doc = Document(word_path)
        
        # 提取所有段落文本
        full_text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text += paragraph.text + "\n"

        return full_text

    except Exception as e:
        print(f"处理Word文件时出错: {e}")
        return None